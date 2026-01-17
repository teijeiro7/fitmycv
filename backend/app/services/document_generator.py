"""
Document generation service for creating optimized CVs in DOCX and PDF formats.
"""
from pathlib import Path
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import tempfile
import os


class DocumentGenerator:
    """Service for generating CV documents from optimized content"""

    def __init__(self):
        self.template_dir = Path(__file__).parent.parent / "templates"

    def generate_docx(
        self,
        optimized_content: Dict[str, Any],
        output_path: Optional[Path] = None,
        template_name: str = "modern"
    ) -> Path:
        """
        Generate a DOCX file from optimized content.

        Args:
            optimized_content: Dictionary with optimized CV sections
            output_path: Path to save the file. If None, generates a temp file
            template_name: Name of the template to use

        Returns:
            Path to the generated file
        """
        doc = Document()

        # Set up document styles
        self._setup_styles(doc)

        # Add sections
        self._add_header_section(doc, optimized_content.get("header", {}))
        self._add_summary_section(doc, optimized_content.get("summary", ""))
        self._add_experience_section(doc, optimized_content.get("experience", ""))
        self._add_projects_section(doc, optimized_content.get("projects", ""))
        self._add_skills_section(doc, optimized_content.get("skills", ""))
        self._add_education_section(doc, optimized_content.get("education", ""))

        # Save the document
        if output_path is None:
            output_path = Path(tempfile.gettempdir()) / f"cv_{id(self)}.docx"

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        return output_path

    def _setup_styles(self, doc: Document):
        """Set up document styles"""
        styles = doc.styles

        # Normal style
        normal = styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)

        # Heading styles
        for i in range(1, 5):
            heading = styles[f'Heading {i}']
            heading.font.name = 'Calibri'
            heading.font.bold = True
            heading.font.size = Pt(16 - i * 2)
            heading.font.color.rgb = RGBColor(0, 0, 0)

    def _add_header_section(self, doc: Document, header_data: Dict):
        """Add the header/contact information section"""
        if not header_data:
            return

        # Name and title
        name = header_data.get("name", "Your Name")
        title = header_data.get("title", "")

        p = doc.add_paragraph()
        run = p.add_run(name)
        run.font.size = Pt(18)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if title:
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(100, 100, 100)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact info
        contact = []
        if header_data.get("email"):
            contact.append(header_data["email"])
        if header_data.get("phone"):
            contact.append(header_data["phone"])
        if header_data.get("location"):
            contact.append(header_data["location"])
        if header_data.get("linkedin"):
            contact.append(header_data["linkedin"])

        if contact:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for i, item in enumerate(contact):
                run = p.add_run(item)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(100, 100, 100)
                if i < len(contact) - 1:
                    p.add_run(" | ")

        doc.add_paragraph()  # Spacer

    def _add_section_heading(self, doc: Document, text: str):
        """Add a section heading with styling"""
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(50, 50, 150)

        # Add a horizontal line effect
        p = doc.add_paragraph()
        p.add_run("_" * 80)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_summary_section(self, doc: Document, summary: str):
        """Add professional summary section"""
        if not summary:
            return

        self._add_section_heading(doc, "Professional Summary")

        p = doc.add_paragraph(summary)
        p.runs[0].font.size = Pt(11)

        doc.add_paragraph()  # Spacer

    def _add_experience_section(self, doc: Document, experience: str):
        """Add work experience section"""
        if not experience:
            return

        self._add_section_heading(doc, "Professional Experience")

        # Parse experience into entries if it's structured
        if isinstance(experience, list):
            for exp in experience:
                self._add_experience_entry(doc, exp)
        else:
            # Treat as raw text
            p = doc.add_paragraph(experience)

        doc.add_paragraph()  # Spacer

    def _add_experience_entry(self, doc: Document, exp: Dict):
        """Add a single experience entry"""
        # Title and company
        title = exp.get("title", "")
        company = exp.get("company", "")
        date = exp.get("date", "")

        header_text = f"{title}"
        if company:
            header_text += f" | {company}"
        if date:
            header_text += f" | {date}"

        p = doc.add_paragraph()
        run = p.add_run(header_text)
        run.font.bold = True
        run.font.size = Pt(11)

        # Achievements/Bullet points
        achievements = exp.get("achievements", [])
        if achievements:
            for achievement in achievements:
                p = doc.add_paragraph(achievement, style='List Bullet')
                p.runs[0].font.size = Pt(10)

        doc.add_paragraph()

    def _add_projects_section(self, doc: Document, projects: Any):
        """Add projects section (including GitHub projects)"""
        if not projects:
            return

        self._add_section_heading(doc, "Projects")

        if isinstance(projects, list):
            for project in projects:
                self._add_project_entry(doc, project)
        elif isinstance(projects, str):
            doc.add_paragraph(projects)

        doc.add_paragraph()

    def _add_project_entry(self, doc: Document, project: Dict):
        """Add a single project entry"""
        name = project.get("name", "Project")
        description = project.get("description", "")
        url = project.get("url", "")
        technologies = project.get("technologies", [])

        # Project name
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.font.bold = True
        run.font.size = Pt(11)

        # Description
        if description:
            p = doc.add_paragraph(description)

        # Technologies
        if technologies:
            tech_text = "Technologies: " + ", ".join(technologies)
            p = doc.add_paragraph(tech_text)
            p.runs[0].font.italic = True
            p.runs[0].font.size = Pt(10)

        # URL
        if url:
            p = doc.add_paragraph()
            p.add_run("Link: ")
            run = p.add_run(url)
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.font.underline = True
            p.runs[0].font.size = Pt(9)

        doc.add_paragraph()

    def _add_skills_section(self, doc: Document, skills: Any):
        """Add skills section"""
        if not skills:
            return

        self._add_section_heading(doc, "Skills")

        if isinstance(skills, list):
            # Group by category if possible, otherwise just list
            p = doc.add_paragraph()
            p.add_run(" | ".join(skills))
        elif isinstance(skills, dict):
            # Categorized skills
            for category, skill_list in skills.items():
                p = doc.add_paragraph()
                run = p.add_run(f"{category}: ")
                run.font.bold = True
                p.add_run(", ".join(skill_list) if isinstance(skill_list, list) else str(skill_list))
        else:
            doc.add_paragraph(str(skills))

        doc.add_paragraph()

    def _add_education_section(self, doc: Document, education: Any):
        """Add education section"""
        if not education:
            return

        self._add_section_heading(doc, "Education")

        if isinstance(education, list):
            for edu in education:
                self._add_education_entry(doc, edu)
        else:
            doc.add_paragraph(education)

    def _add_education_entry(self, doc: Document, edu: Dict):
        """Add a single education entry"""
        degree = edu.get("degree", "")
        school = edu.get("school", "")
        year = edu.get("year", "")

        text = f"{degree}"
        if school:
            text += f" - {school}"
        if year:
            text += f", {year}"

        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)

        doc.add_paragraph()

    def generate_pdf(
        self,
        docx_path: Optional[Path] = None,
        optimized_content: Optional[Dict] = None,
        output_path: Optional[Path] = None
    ) -> Optional[Path]:
        """
        Generate a PDF file.

        Args:
            docx_path: Path to source DOCX file (if available)
            optimized_content: Content to generate from (if no DOCX)
            output_path: Path to save the PDF

        Returns:
            Path to generated PDF or None if generation failed
        """
        if output_path is None:
            output_path = Path(tempfile.gettempdir()) / f"cv_{id(self)}.pdf"

        output_path.parent.mkdir(parents=True, exist_ok=True)

        # First, ensure we have a DOCX file
        if docx_path is None and optimized_content:
            docx_path = self.generate_docx(optimized_content)

        if docx_path is None or not Path(docx_path).exists():
            return None

        # Try to convert using LibreOffice (most reliable)
        result = self._convert_with_libreoffice(docx_path, output_path)
        if result:
            return result

        # Fallback: try using docx2pdf if available
        try:
            from docx2pdf import convert
            convert(str(docx_path), str(output_path))
            return output_path
        except ImportError:
            pass
        except Exception:
            pass

        # If all else fails, return None (PDF generation is optional)
        return None

    def _convert_with_libreoffice(self, docx_path: Path, output_path: Path) -> Optional[Path]:
        """Convert DOCX to PDF using LibreOffice"""
        try:
            # Common LibreOffice installation paths
            libreoffice_paths = [
                "/usr/bin/libreoffice",
                "/usr/local/bin/libreoffice",
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
                "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
            ]

            soffice = None
            for path in libreoffice_paths:
                if Path(path).exists():
                    soffice = path
                    break

            if not soffice:
                return None

            # Convert using LibreOffice headless mode
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf",
                 "--outdir", str(output_path.parent), str(docx_path)],
                capture_output=True,
                timeout=30
            )

            # LibreOffice creates PDF with same basename
            expected_pdf = output_path.parent / f"{docx_path.stem}.pdf"
            if expected_pdf.exists():
                if expected_pdf != output_path:
                    expected_pdf.rename(output_path)
                return output_path

        except (subprocess.TimeoutExpired, FileNotFoundError, Exception):
            pass

        return None


def format_gherkin_bullet(text: str) -> str:
    """Format text as a bullet point using Gherkin syntax for clarity"""
    return f"• {text}"


def clean_text_for_docx(text: str) -> str:
    """Clean text for DOCX generation"""
    if not text:
        return ""
    # Remove problematic characters
    text = text.replace("\x00", "").replace("\x08", "")
    return text.strip()
