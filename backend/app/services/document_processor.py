"""
Document processing service for parsing CVs from PDF and DOCX formats.
"""
import io
from pathlib import Path
from typing import Optional, Tuple
from docx import Document
import PyPDF2
import pdfplumber


class DocumentProcessor:
    """Service for processing various document formats"""

    @staticmethod
    def extract_text_from_docx(file_path: Path) -> str:
        """
        Extract text from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text as string
        """
        try:
            doc = Document(file_path)
            text_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n".join(text_parts)

        except Exception as e:
            raise ValueError(f"Error processing DOCX file: {str(e)}")

    @staticmethod
    def extract_text_from_pdf(file_path: Path) -> str:
        """
        Extract text from a PDF file using pdfplumber for better accuracy.

        Args:
            file_path: Path to the PDF file

        Returns:
            Extracted text as string
        """
        try:
            # Try pdfplumber first (better layout preservation)
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text)

                if text_parts:
                    return "\n".join(text_parts)

        except Exception as e:
            # Fall back to PyPDF2 if pdfplumber fails
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_parts = []

                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(page_text)

                    return "\n".join(text_parts)

            except Exception as e2:
                raise ValueError(f"Error processing PDF file: {str(e)}, fallback also failed: {str(e2)}")

    @staticmethod
    def extract_text(file_path: Path) -> str:
        """
        Extract text from a supported document file (PDF or DOCX).

        Args:
            file_path: Path to the document file

        Returns:
            Extracted text as string

        Raises:
            ValueError: If file format is not supported or extraction fails
        """
        file_extension = file_path.suffix.lower()

        if file_extension == '.docx':
            return DocumentProcessor.extract_text_from_docx(file_path)
        elif file_extension == '.pdf':
            return DocumentProcessor.extract_text_from_pdf(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: .pdf, .docx")

    @staticmethod
    def get_document_metadata(file_path: Path) -> dict:
        """
        Extract metadata from a document file.

        Args:
            file_path: Path to the document file

        Returns:
            Dictionary with metadata (page count, word count, etc.)
        """
        file_extension = file_path.suffix.lower()
        text = DocumentProcessor.extract_text(file_path)
        word_count = len(text.split()) if text else 0
        char_count = len(text) if text else 0

        metadata = {
            "file_name": file_path.name,
            "file_extension": file_extension,
            "word_count": word_count,
            "character_count": char_count,
        }

        if file_extension == '.pdf':
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata["page_count"] = len(pdf_reader.pages)
            except:
                metadata["page_count"] = 0

        elif file_extension == '.docx':
            try:
                doc = Document(file_path)
                # Approximate page count (doesn't account for formatting)
                metadata["page_count"] = max(1, len(doc.paragraphs) // 25)
            except:
                metadata["page_count"] = 0

        return metadata

    @staticmethod
    def validate_file_extension(filename: str, allowed_extensions: list[str]) -> Tuple[bool, Optional[str]]:
        """
        Validate if a file has an allowed extension.

        Args:
            filename: Name of the file to validate
            allowed_extensions: List of allowed extensions (e.g., ['.pdf', '.docx'])

        Returns:
            Tuple of (is_valid, actual_extension)
        """
        file_extension = Path(filename).suffix.lower()
        is_valid = file_extension in allowed_extensions
        return is_valid, file_extension


def parse_resume_structure(text: str) -> dict:
    """
    Attempt to parse a resume into structured sections.
    This is a simple heuristic-based parser.

    Args:
        text: Raw resume text

    Returns:
        Dictionary with parsed sections
    """
    lines = text.split('\n')
    sections = {
        "header": "",
        "summary": "",
        "experience": "",
        "education": "",
        "skills": "",
        "projects": "",
        "languages": "",
        "certifications": "",
        "other": ""
    }

    current_section = "header"
    current_content = []

    # Common section headers (case-insensitive)
    section_keywords = {
        "summary": ["summary", "profile", "about", "objective", "professional summary"],
        "experience": ["experience", "work experience", "employment", "work history", "professional experience"],
        "education": ["education", "academic", "qualifications", "academic background"],
        "skills": ["skills", "technical skills", "competencies", "expertise", "technologies"],
        "projects": ["projects", "portfolio", "key projects"],
        "languages": ["languages", "language proficiency"],
        "certifications": ["certifications", "certificates", "credentials"]
    }

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        # Check if this line is a section header
        is_section_header = False
        line_lower = line_stripped.lower()

        for section, keywords in section_keywords.items():
            if any(keyword in line_lower for keyword in keywords):
                # Save previous section
                if current_content:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = section
                current_content = []
                is_section_header = True
                break

        if not is_section_header:
            current_content.append(line_stripped)

    # Don't forget the last section
    if current_content:
        sections[current_section] = "\n".join(current_content).strip()

    return sections
